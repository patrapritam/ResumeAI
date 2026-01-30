"""
Text Extractor Module
Extract text content from PDF and DOCX files
"""

from PyPDF2 import PdfReader
from docx import Document
import re


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text content
    """
    try:
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        full_text = '\n'.join(text_parts)
        
        # Clean up the text
        full_text = clean_text(full_text)
        
        return full_text
    
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        full_text = '\n'.join(text_parts)
        
        # Clean up the text
        full_text = clean_text(full_text)
        
        return full_text
    
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep common punctuation
    text = re.sub(r'[^\w\s\.\,\;\:\-\@\#\+\/\(\)]', '', text)
    
    # Normalize line breaks
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()
